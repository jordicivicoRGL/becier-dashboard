# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def add_image_box(doc, label):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(label)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    shade_cell(cell, 'EEEEEE')
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom']:
        m = OxmlElement('w:' + side)
        m.set(qn('w:w'), '180')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tc_pr.append(tcMar)
    doc.add_paragraph()


def add_cta_button(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.bold = True
    shade_cell(cell, 'FFFFFF')
    tbl = table._tbl
    tblPr_el = tbl.find(qn('w:tblPr'))
    if tblPr_el is not None:
        existing_w = tblPr_el.find(qn('w:tblW'))
        if existing_w is not None:
            tblPr_el.remove(existing_w)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '3200')
        tblW.set(qn('w:type'), 'dxa')
        tblPr_el.append(tblW)
    doc.add_paragraph()


def add_bold_label(doc, label, value, underline=False):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    if underline:
        r1.underline = True
    r2 = p.add_run(value)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    return p


# ─── ENCABEZADO ───────────────────────────────────────────────
add_bold_label(doc, 'ASUNTO: ', 'SS26: tejido que trabaja contigo')
add_bold_label(doc, 'PREENCABEZADO: ', 'Colección completa disponible. Ultraligero, sin costuras, badana 3D.', underline=True)
doc.add_paragraph()

# ─── HERO ─────────────────────────────────────────────────────
add_image_box(
    doc,
    'IMAGEN: GIF o foto hero — ciclista en ruta con kit SS26 completo, plano lateral a alta velocidad\n'
    'Referencia Pinterest: pinterest.com/search/pins/?q=cycling+kit+road+action+photography'
)

p = doc.add_paragraph()
r = p.add_run('Tejido técnico. Confección sin dobladillos. SS26 ya está aquí.')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(13)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run(
    'Diseñada para temperaturas entre 20-35°C y entrenamientos de alta intensidad. '
    'Confección bonding sin dobladillos en puntos clave, tejidos de secado rápido '
    'con transpirabilidad máxima y badanas 3D moldeadas en tres capas.'
)
r.font.name = 'Arial'
r.font.size = Pt(11)
doc.add_paragraph()

# ─── MAILLOTS ─────────────────────────────────────────────────
add_image_box(
    doc,
    'IMAGEN: Maillot en ruta — detalle del tejido en movimiento, plano lateral o tres cuartos\n'
    'Referencia Pinterest: pinterest.com/search/pins/?q=cycling+jersey+technical+fabric+road'
)

p = doc.add_paragraph()
r = p.add_run('MAILLOTS SS26')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(12)
doc.add_paragraph()

maillots = [
    ('Aero Light SS · €125', 'Tejido ultraligero, UPF 50+, cremallera con bloqueo frontal ultra suave. Transpirabilidad nivel 5. Para 25-35°C.'),
    ('Overline · €135', 'Estructura de semi-malla (mid-mesh). Secado de humedad en marcha. Para calor extremo y humedad alta.'),
    ('Performance SS · €130', 'Paneles de ventilación interior y bolsillos con estructura de panel de abeja. Tejido de punto circular.'),
    ('Air 2.0 LE · €125', 'Estructura semiabierta, confección bonding, bolsillos lumbares minimalistas.'),
]
for name, desc in maillots:
    p = doc.add_paragraph()
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(' — ' + desc)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)

doc.add_paragraph()
add_cta_button(doc, 'VER MAILLOTS SS26 →')

# ─── CULOTES ──────────────────────────────────────────────────
add_image_box(
    doc,
    'IMAGEN: Culottes en ruta — plano medio-bajo del pedaleo, pernera visible\n'
    'Referencia Pinterest: pinterest.com/search/pins/?q=cycling+bib+shorts+road+movement'
)

p = doc.add_paragraph()
r = p.add_run('CULOTTES SS26')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(12)
doc.add_paragraph()

culotes = [
    ('Air 2.0 · €170', 'Pernera y tirantes cortados con láser, sin costuras. Badana de densidad múltiple, 3 capas moldeadas en 3D. Intensidad super alta.'),
    ('Cargo · €160', 'Badana 3D + bolsillos laterales integrados (móvil, geles, herramientas). Tejido certificado OEKO-TEX®. Para rutas largas.'),
    ('Performance · €150', 'Tejido italiano certificado OEKO-TEX®, badana 3D multicapa, suspensores cortados con láser, sin costuras. Hasta 7 horas.'),
]
for name, desc in culotes:
    p = doc.add_paragraph()
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(' — ' + desc)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)

doc.add_paragraph()
add_cta_button(doc, 'VER CULOTTES SS26 →')

# ─── CIERRE ───────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('La SS26 ya está en ruta. Encuentra tu talla antes de que se agote.')
r.font.name = 'Arial'
r.font.size = Pt(11)
doc.add_paragraph()
add_cta_button(doc, 'VER COLECCIÓN COMPLETA SS26 →')

# ─── REFERENCIAS PINTEREST ────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('REFERENCIAS VISUALES (Pinterest)')
r.bold = True
r.font.name = 'Arial'
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

refs = [
    ('Hero / acción en ruta:', 'pinterest.com/search/pins/?q=cycling+kit+road+action+photography'),
    ('Detalle tejido técnico:', 'pinterest.com/search/pins/?q=technical+cycling+fabric+closeup'),
    ('Culotes en movimiento:', 'pinterest.com/search/pins/?q=cycling+bib+shorts+road+movement'),
]
for label, url in refs:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    r2 = p.add_run(url)
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)

os.makedirs('outputs', exist_ok=True)
path = 'outputs/mailing_ss26_twentyone.docx'
doc.save(path)
print('Guardado: ' + path)
