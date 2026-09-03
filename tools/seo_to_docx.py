"""
Convierte el informe SEO de Diagonal CQ (formato Markdown) a Word (.docx).
Uso: python -m tools.seo_to_docx <ruta_md> [ruta_docx]
Si no se indica ruta_docx, se guarda en el mismo directorio con extension .docx.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COLOR_SLIDE   = RGBColor(0xC0, 0x39, 0x2B)   # rojo oscuro para cabeceras de slide
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x2E)   # azul oscuro para titulos internos


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_run_formatted(paragraph, text: str):
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def _parse_table(doc: Document, lines: list, start: int):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i])
        i += 1

    data_lines = [l for l in table_lines if not re.match(r"^\s*\|[-:| ]+\|\s*$", l)]
    if not data_lines:
        return i

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return i

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.clear()
            _add_run_formatted(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                _set_cell_bg(cell, "1A1A2E")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()
    return i


def convert_md_to_docx(md_path: str, docx_path: str):
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        # Separadores (--- o lineas de guiones)
        if re.match(r"^[-]+$", stripped):
            i += 1
            continue

        # Cabecera de SLIDE — detecta tanto "SLIDE: X" solo como con barras
        slide_match = re.match(r"^[━\-]*\s*SLIDE:\s*(.+?)\s*[━\-]*$", stripped)
        if slide_match:
            slide_title = slide_match.group(1)
            p = doc.add_paragraph()
            run = p.add_run("SLIDE: " + slide_title)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_SLIDE
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Encabezados Markdown (# ## ###)
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            font_size = {1: 18, 2: 15, 3: 13, 4: 12}.get(level, 12)
            run.font.size = Pt(font_size)
            run.font.color.rgb = COLOR_HEADING
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Tabla Markdown
        if stripped.startswith("|"):
            i = _parse_table(doc, lines, i)
            continue

        # Blockquote (> texto)
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Bullet (- texto o * texto)
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_run_formatted(p, text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Lista numerada
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_run_formatted(p, text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Parrafo normal
        p = doc.add_paragraph()
        _add_run_formatted(p, stripped)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python -m tools.seo_to_docx <ruta_md> [ruta_docx]")
        sys.exit(1)

    md = sys.argv[1]
    docx = sys.argv[2] if len(sys.argv) > 2 else str(Path(md).with_suffix(".docx"))
    result = convert_md_to_docx(md, docx)
    print(f"Word guardado en: {result}")
